"""Regenerate the test fixtures in this directory.

Run with::

    uv run python tests/fixtures/generate_fixtures.py

The generated files are committed so the test suite is hermetic and needs no
generation step. This script exists so the fixtures stay reproducible and
auditable — a reviewer can see exactly what "a legacy-font PDF" means here.

Provenance of the Sinhala sample strings is recorded in ``README.md`` beside
this script.
"""

from __future__ import annotations

import pathlib
import sys

FIXTURES = pathlib.Path(__file__).parent

# --- Sample content -------------------------------------------------------

# Real FM-Abhaya legacy bytes. In the legacy encoding each Latin-1 codepoint is
# a glyph index, not a letter; pandukabhaya maps these to Sinhala Unicode.
# Converts to: අධ්‍යාපන ප්‍රකාශන දෙපාර්තමේන්තුව
LEGACY_SINHALA = "wOHdmk m%ldYk fomd¾;fïka;=j"

# The ASCII control string. Blanket legacy conversion turns this into
# "අඅඅගැාමචමඉගටදඩගකන" — the prototype bug this whole design exists to prevent.
ASCII_CONTROL = "www.edupub.gov.lk"

# Genuine Sinhala Unicode (U+0D80-U+0DFF). Must survive normalisation untouched.
UNICODE_SINHALA = "අධ්‍යාපන ප්‍රකාශන දෙපාර්තමේන්තුව"
UNICODE_SINHALA_2 = "සියලු ම පෙළපොත් නොමිලේ බෙදා දෙනු ලැබේ"

SINHALA_FONT_FILE = pathlib.Path("C:/Windows/Fonts/iskpota.ttf")


# --- sample_unicode.pdf ---------------------------------------------------


def build_unicode_pdf(path: pathlib.Path) -> str:
    """Born-digital PDF carrying real Sinhala Unicode in an embedded font.

    Returns a short description of the route taken, for the provenance record.
    """
    import pymupdf

    doc = pymupdf.open()
    page = doc.new_page(width=612, height=792)

    if SINHALA_FONT_FILE.exists():
        page.insert_font(fontname="iskpota", fontfile=str(SINHALA_FONT_FILE))
        route = f"embedded {SINHALA_FONT_FILE.name}"
        font = "iskpota"
    else:  # pragma: no cover - depends on the host machine
        route = "FALLBACK: no system Sinhala font found, used Helvetica"
        font = "helv"

    page.insert_text((72, 700), UNICODE_SINHALA, fontname=font, fontsize=18)
    page.insert_text((72, 660), UNICODE_SINHALA_2, fontname=font, fontsize=14)
    page.insert_text((72, 620), ASCII_CONTROL, fontname="helv", fontsize=12)
    # Without subsetting, the whole 538 KB Iskoola Pota file is embedded.
    doc.subset_fonts()
    doc.save(path, garbage=4, deflate=True)
    doc.close()
    return route


# --- sample_legacy_font.pdf -----------------------------------------------


def build_legacy_pdf(path: pathlib.Path) -> str:
    """Hand-authored PDF with a *non-embedded* ``/BaseFont /FMAbhaya``.

    We do not ship an FM Abhaya TTF (proprietary, and none is available). We do
    not need one: the two things the code under test consumes are the span's
    font *name* and the span's raw *bytes*, and PyMuPDF surfaces both correctly
    for a non-embedded font. Verified during planning:
    ``span["font"] == "FMAbhaya"`` and the latin-1 payload arrives intact.

    A second span in Helvetica carries ``ASCII_CONTROL`` so the
    ASCII-preservation regression has a target in the same document.
    """
    content = (
        f"BT /F1 18 Tf 72 700 Td ({LEGACY_SINHALA}) Tj ET\n"
        f"BT /F2 12 Tf 72 660 Td ({ASCII_CONTROL}) Tj ET\n"
    ).encode("latin-1")

    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << /Font << /F1 5 0 R /F2 6 0 R >> >> /Contents 4 0 R >>",
        b"<< /Length "
        + str(len(content)).encode()
        + b" >>\nstream\n"
        + content
        + b"\nendstream",
        b"<< /Type /Font /Subtype /TrueType /BaseFont /FMAbhaya "
        b"/Encoding /WinAnsiEncoding >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica "
        b"/Encoding /WinAnsiEncoding >>",
    ]
    path.write_bytes(_assemble_pdf(objects))
    return "hand-authored, non-embedded /BaseFont /FMAbhaya"


def _assemble_pdf(objects: list[bytes]) -> bytes:
    """Serialise numbered PDF objects into a minimal valid PDF with an xref."""
    out = bytearray(b"%PDF-1.4\n")
    offsets: list[int] = []
    for number, body in enumerate(objects, start=1):
        offsets.append(len(out))
        out += f"{number} 0 obj\n".encode() + body + b"\nendobj\n"

    xref_at = len(out)
    out += f"xref\n0 {len(objects) + 1}\n".encode()
    out += b"0000000000 65535 f \n"
    for offset in offsets:
        out += f"{offset:010d} 00000 n \n".encode()
    out += (
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
        f"startxref\n{xref_at}\n%%EOF\n"
    ).encode()
    return bytes(out)


# --- sample_scanned.pdf / sample_mixed.pdf --------------------------------


def _rasterise_page(source_page, target_doc, dpi: int = 150) -> None:
    """Append an image-only copy of ``source_page`` to ``target_doc``.

    Greyscale + JPEG keeps the committed fixture around 100 KB instead of the
    ~6 MB an RGB pixmap produces, while staying legible enough at 150 dpi for
    Tesseract to work with.
    """
    import pymupdf

    pixmap = source_page.get_pixmap(dpi=dpi, colorspace=pymupdf.csGRAY)
    page = target_doc.new_page(
        width=source_page.rect.width, height=source_page.rect.height
    )
    page.insert_image(page.rect, stream=pixmap.tobytes("jpeg", jpg_quality=70))


def build_scanned_pdf(source: pathlib.Path, path: pathlib.Path) -> str:
    """Image-only PDF: page 0 of ``source`` rasterised, with no text layer."""
    import pymupdf

    src = pymupdf.open(source)
    doc = pymupdf.open()
    _rasterise_page(src[0], doc)
    doc.save(path)
    doc.close()
    src.close()
    return "page 0 of sample_unicode.pdf rasterised at 150 dpi"


def build_mixed_pdf(source: pathlib.Path, path: pathlib.Path) -> str:
    """Page 0 born-digital, page 1 rasterised — exercises per-page OCR routing."""
    import pymupdf

    src = pymupdf.open(source)
    doc = pymupdf.open()
    doc.insert_pdf(src, from_page=0, to_page=0)
    _rasterise_page(src[0], doc)
    doc.save(path)
    doc.close()
    src.close()
    return "page 0 born-digital, page 1 rasterised"


# --- sample.docx ----------------------------------------------------------


def build_docx(path: pathlib.Path) -> str:
    """Word document mixing legacy and Unicode runs inside one paragraph.

    Spec Section 9 requires the same paragraph to carry both, so per-run
    detection can be distinguished from per-paragraph detection. A third run
    carries no explicit font, to exercise the style-inheritance chain.
    """
    from docx import Document

    document = Document()

    paragraph = document.add_paragraph()
    legacy_run = paragraph.add_run(LEGACY_SINHALA)
    legacy_run.font.name = "FMAbhaya"
    unicode_run = paragraph.add_run(" " + UNICODE_SINHALA)
    unicode_run.font.name = "Iskoola Pota"
    paragraph.add_run(" " + ASCII_CONTROL)  # no explicit font: inherits Normal

    document.add_paragraph("Between the paragraph and the table.")

    table = document.add_table(rows=2, cols=2)
    cells = [
        (0, 0, LEGACY_SINHALA, "FMAbhaya"),
        (0, 1, UNICODE_SINHALA, "Iskoola Pota"),
        (1, 0, LEGACY_SINHALA, "FMBindumathix"),
        (1, 1, ASCII_CONTROL, None),
    ]
    for row, col, text, font_name in cells:
        run = table.cell(row, col).paragraphs[0].add_run(text)
        if font_name is not None:
            run.font.name = font_name

    document.add_paragraph("After the table.")
    document.save(path)
    return "3 paragraphs + 2x2 table, mixed runs"


# --- sample.xlsx ----------------------------------------------------------


def build_xlsx(path: pathlib.Path) -> str:
    """Workbook with two sheets, legacy-font cells, a formula and empty cells."""
    from openpyxl import Workbook
    from openpyxl.styles import Font

    workbook = Workbook()

    legacy = workbook.active
    legacy.title = "Legacy"
    legacy["A1"] = LEGACY_SINHALA
    legacy["A1"].font = Font(name="FMAbhaya")
    legacy["A2"] = LEGACY_SINHALA
    legacy["A2"].font = Font(name="FMBindumathix")
    legacy["B1"] = ASCII_CONTROL  # default font: None
    legacy["B3"] = 42
    legacy["C3"] = "=B3*2"  # formula; data_only reads the cached value

    unicode_sheet = workbook.create_sheet("Unicode")
    unicode_sheet["A1"] = UNICODE_SINHALA
    unicode_sheet["A1"].font = Font(name="Nirmala UI")
    unicode_sheet["A3"] = UNICODE_SINHALA_2  # A2 left empty on purpose
    unicode_sheet["A3"].font = Font(name="Iskoola Pota")

    workbook.save(path)
    return "2 sheets: Legacy (FM fonts, formula), Unicode (Nirmala UI/Iskoola Pota)"


# --- driver ---------------------------------------------------------------


def main() -> int:
    FIXTURES.mkdir(parents=True, exist_ok=True)
    routes: dict[str, str] = {}

    unicode_pdf = FIXTURES / "sample_unicode.pdf"
    routes["sample_unicode.pdf"] = build_unicode_pdf(unicode_pdf)
    routes["sample_legacy_font.pdf"] = build_legacy_pdf(
        FIXTURES / "sample_legacy_font.pdf"
    )
    routes["sample_scanned.pdf"] = build_scanned_pdf(
        unicode_pdf, FIXTURES / "sample_scanned.pdf"
    )
    routes["sample_mixed.pdf"] = build_mixed_pdf(
        unicode_pdf, FIXTURES / "sample_mixed.pdf"
    )
    routes["sample.docx"] = build_docx(FIXTURES / "sample.docx")
    routes["sample.xlsx"] = build_xlsx(FIXTURES / "sample.xlsx")

    for name, route in routes.items():
        size = (FIXTURES / name).stat().st_size
        print(f"  {name:26} {size:>8,} bytes   {route}")

    _self_check()
    print("\nAll fixtures generated and self-checked.")
    return 0


def _self_check() -> None:
    """Assert the structural properties the test suite depends on."""
    import pymupdf

    legacy = pymupdf.open(FIXTURES / "sample_legacy_font.pdf")
    fonts = {
        span["font"]
        for block in legacy[0].get_text("dict")["blocks"]
        for line in block.get("lines", [])
        for span in line["spans"]
    }
    assert "FMAbhaya" in fonts, f"legacy fixture lost its font name: {fonts}"
    assert ASCII_CONTROL in legacy[0].get_text(), "legacy fixture lost its ASCII span"
    legacy.close()

    scanned = pymupdf.open(FIXTURES / "sample_scanned.pdf")
    assert not scanned[0].get_text().strip(), "scanned fixture has a text layer"
    coverage = _largest_image_coverage(scanned[0])
    assert coverage > 0.95, f"scanned fixture image covers only {coverage:.2%}"
    scanned.close()

    mixed = pymupdf.open(FIXTURES / "sample_mixed.pdf")
    assert len(mixed) == 2, "mixed fixture should have exactly 2 pages"
    assert mixed[0].get_text().strip(), "mixed page 0 should be born-digital"
    assert not mixed[1].get_text().strip(), "mixed page 1 should be image-only"
    mixed.close()


def _largest_image_coverage(page) -> float:
    page_area = abs(page.rect)
    if not page_area:
        return 0.0
    return max(
        (abs(pymupdf_rect(info["bbox"]) & page.rect) / page_area
         for info in page.get_image_info()),
        default=0.0,
    )


def pymupdf_rect(bbox):
    import pymupdf

    return pymupdf.Rect(bbox)


if __name__ == "__main__":
    sys.exit(main())
