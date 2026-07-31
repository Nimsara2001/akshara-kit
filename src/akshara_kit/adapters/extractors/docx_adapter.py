"""DOCX extraction via python-docx.

One deterministic path, so this adapter is also the DOCX coordinator. Font
detection is per **run**, not per paragraph: a single paragraph routinely mixes
a legacy-font run with a Unicode one, and per-paragraph detection would either
miss the legacy text or corrupt the Unicode text.

Headers, footers, footnotes and text boxes are out of scope for this build.
"""

from __future__ import annotations

import time
from typing import TYPE_CHECKING, Iterator

from akshara_kit.contracts.extraction import ExtractionResult, SourceFormat
from akshara_kit.eye.errors import AdapterUnavailableError, ExtractionFailedError
from akshara_kit.eye.quality_probe import score

if TYPE_CHECKING:
    from docx.document import Document as DocumentType
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    from akshara_kit.eye.font_detection import SpanFont

__all__ = ["BACKEND_ID", "extract"]

BACKEND_ID = "python-docx"

_CELL_SEPARATOR = "\t"
_BLOCK_SEPARATOR = "\n"


def extract(file_path: str) -> ExtractionResult:
    """Extract every run, including table cells, in document order.

    Realises Section 5's DOCX path and Section 6.3's per-run font detection.
    """
    from akshara_kit.eye.encoding_normaliser import normalise_spans

    started = time.perf_counter()
    document = _open(file_path)
    spans = list(_iter_run_spans(document))

    outcome = normalise_spans(spans)
    return ExtractionResult(
        text=outcome.text,
        backend_id=BACKEND_ID,
        source_format=SourceFormat.DOCX,
        latency_seconds=time.perf_counter() - started,
        quality=score(outcome.text),
        font_detection_method=outcome.method,
        detected_legacy_fonts=outcome.converted_fonts,
        unmapped_legacy_fonts=outcome.unmapped_fonts,
        metadata={"runs": sum(1 for span in spans if not span.is_separator)},
    )


def _open(file_path: str) -> DocumentType:
    """Open the document, translating library failures into named errors."""
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - depends on install extras
        raise AdapterUnavailableError(
            "python-docx is not installed; install the 'docx' extra"
        ) from exc

    try:
        return Document(file_path)
    except Exception as exc:
        raise ExtractionFailedError(
            f"{BACKEND_ID} could not open {file_path}: {exc}"
        ) from exc


def _iter_block_items(parent) -> Iterator[Paragraph | Table]:
    """Yield paragraphs and tables in true document order.

    ``document.paragraphs`` and ``document.tables`` are separate collections,
    so reading them in turn loses the interleaving. Walking the underlying XML
    body is the only way to keep a table in its real position.
    """
    from docx.document import Document as DocumentType
    from docx.oxml.ns import qn
    from docx.table import Table, _Cell
    from docx.text.paragraph import Paragraph

    if isinstance(parent, DocumentType):
        element = parent.element.body
    elif isinstance(parent, _Cell):
        element = parent._tc
    else:  # pragma: no cover - defensive
        raise TypeError(f"Cannot iterate blocks of {type(parent)!r}")

    for child in element.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def _iter_run_spans(document: DocumentType) -> Iterator[SpanFont]:
    """Yield one span per run across the whole document, in reading order."""
    from docx.table import Table

    default_font = _default_font(document)
    for block in _iter_block_items(document):
        if isinstance(block, Table):
            yield from _iter_table_spans(block, default_font)
        else:
            yield from _iter_paragraph_spans(block, default_font)
        yield _separator(_BLOCK_SEPARATOR)


def _iter_table_spans(table: Table, default_font: str) -> Iterator[SpanFont]:
    """Yield spans for every cell, recursing into nested tables."""
    from docx.table import Table as TableType

    for row in table.rows:
        for cell_index, cell in enumerate(row.cells):
            if cell_index:
                yield _separator(_CELL_SEPARATOR)
            for block in _iter_block_items(cell):
                if isinstance(block, TableType):
                    yield from _iter_table_spans(block, default_font)
                else:
                    yield from _iter_paragraph_spans(block, default_font)
        yield _separator(_BLOCK_SEPARATOR)


def _iter_paragraph_spans(paragraph: Paragraph, default_font: str) -> Iterator[SpanFont]:
    """Yield one span per run in a paragraph."""
    from akshara_kit.eye.font_detection import SpanFont

    for run in paragraph.runs:
        if not run.text:
            continue
        yield SpanFont(text=run.text, font=_effective_font(run, default_font))


def _effective_font(run, default_font: str) -> str:
    """Resolve a run's font through the style inheritance chain.

    ``run.font.name`` is ``None`` whenever the font comes from a style, which
    is the common case in real documents; treating that as "no font" would miss
    most legacy text.
    """
    candidates = (
        run.font.name,
        getattr(getattr(run, "style", None), "font", None)
        and run.style.font.name,
        default_font,
    )
    return next((name for name in candidates if name), "")


def _default_font(document: DocumentType) -> str:
    """The document's default run font, from the Normal style.

    Resolved once per extraction and threaded through the iterators, since
    every run that specifies no font of its own inherits it.
    """
    try:
        return document.styles["Normal"].font.name or ""
    except (KeyError, AttributeError):  # pragma: no cover - malformed styles
        return ""


def _separator(text: str) -> SpanFont:
    from akshara_kit.eye.font_detection import separator

    return separator(text)
